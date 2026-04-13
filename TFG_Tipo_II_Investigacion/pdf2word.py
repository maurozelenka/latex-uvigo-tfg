"""
Script de conversion PDF a Word HD para TFG UVigo.
Convierte el PDF a Word y luego inyecta las imagenes de la portada
que pdf2docx no detecta correctamente.
"""
import sys
import os
import shutil
import fitz  # PyMuPDF
import re

def convert_pdf_to_docx(pdf_path, docx_path):
    """Convierte PDF a Word con alta fidelidad visual."""
    from pdf2docx import Converter

    if not os.path.exists(pdf_path):
        print(f"[ERROR] No se encontro el archivo: {pdf_path}")
        sys.exit(1)

    # Usar nombre temporal para evitar conflictos con Word abierto
    temp_docx = docx_path + ".tmp.docx"

    print(f"[INFO] Convirtiendo {pdf_path} -> {docx_path}")

    # Extraer imagenes del PDF (pagina 1)
    doc = fitz.open(pdf_path)
    page = doc[0]
    images_info = page.get_images(full=True)
    extracted_images = []
    
    for img_index, img in enumerate(images_info):
        xref = img[0]
        base_image = doc.extract_image(xref)
        image_bytes = base_image["image"]
        image_ext = base_image["ext"]
        img_filename = f"_temp_img_{img_index}.{image_ext}"
        with open(img_filename, "wb") as f:
            f.write(image_bytes)
        extracted_images.append(img_filename)
        print(f"[INFO] Imagen extraida: {img_filename} ({len(image_bytes)} bytes)")
    
    doc.close()

    try:

        # Conversion principal con pdf2docx
        if os.path.exists(temp_docx):
            os.remove(temp_docx)
        
        cv = Converter(pdf_path)
        cv.convert(temp_docx)
        cv.close()
        print(f"[INFO] Conversion base completada")

        # Post-procesado avanzado
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
        
        wordoc = Document(temp_docx)
            
        # 1. Corregir la Portada (Icono y Tabla)
        target_table = None
        if wordoc.tables:
            for t in wordoc.tables:
                text_in_table = " ".join([c.text for row in t.rows for c in row.cells])
                if "Traballo de Fin de Grao" in text_in_table or "Trabajo de Fin de Grado" in text_in_table:
                    target_table = t
                    break

        if target_table:
            target_table.allow_autofit = False
            for row in target_table.rows:
                row.cells[0].width = Inches(1.2)
                row.cells[1].width = Inches(4.8)

            cell_img = target_table.cell(0, 0)
            # Priorizar imágenes fuente originales sobre las extraídas del PDF
            # para evitar dependencia del orden de extracción (Bug #3)
            icono_path = os.path.join("src", "images", "logo.png")
            if not os.path.exists(icono_path):
                # Fallback: buscar cualquier imagen con nombre "logo" o "icono" en src/images/
                img_dir = os.path.join("src", "images")
                if os.path.isdir(img_dir):
                    for f in os.listdir(img_dir):
                        if f.lower().startswith("logo") or f.lower().startswith("icono"):
                            icono_path = os.path.join(img_dir, f)
                            break
            if not os.path.exists(icono_path) and extracted_images:
                # Último recurso: usar la última imagen extraída del PDF
                icono_path = extracted_images[-1]
                print(f"[WARN] No se encontró logo en src/images/, usando imagen extraída: {icono_path}")
            
            # Insertar imagen en la celda (si existe la ruta o se recuperó)
            if os.path.exists(icono_path):
                # Limpiar cualquier placeholder de imagen rota dejado por pdf2docx
                for p in cell_img.paragraphs:
                    p.clear()
                    
                paragraph_img = cell_img.paragraphs[0] if cell_img.paragraphs else cell_img.add_paragraph()
                paragraph_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run_img = paragraph_img.add_run()
                run_img.add_picture(icono_path, width=Inches(1.1))
                
            # Corregir colisiones y negritas en la celda derecha de la portada
            cell_text = target_table.cell(0, 1)
            labels_negrita = [
                "Traballo de Fin de Grao", "Trabajo de Fin de Grado", 
                "Titor/a:", "Tutor/a:", 
                "Área de coñecemento:", "Área de conocimiento:", 
                "Departamento:"
            ]
            for p in cell_text.paragraphs:
                if "Departamento:" in p.text and "Inform" in p.text and "\nDepartamento:" not in p.text:
                    for run in p.runs:
                        if " Departamento:" in run.text:
                            run.text = run.text.replace(" Departamento:", "\nDepartamento:")
                for run in p.runs:
                    for label in labels_negrita:
                        if label in run.text: run.bold = True

        # 2. Corregir el Indice (TOC)
        print(f"[INFO] Optimizando Indice...")
        for p in wordoc.paragraphs:
            if re.search(r'\d+$', p.text.strip()) and len(p.text) < 150:
                original_has_dots = ".." in p.text
                for run in p.runs:
                    run.text = re.sub(r'^(\d+\.)([^\s])', r'\1 \2', run.text)
                    if re.match(r'^\d+\.$', run.text.strip()): run.text = run.text.strip() + " "
                    if ".." in run.text: run.text = re.sub(r'\.{2,}', '', run.text)
                if "\t" not in p.text:
                    for run in reversed(p.runs):
                        if re.search(r'\d+$', run.text):
                            run.text = re.sub(r'(\s*)(\d+)$', r'\t\2', run.text)
                            break
                leader = WD_TAB_LEADER.DOTS if original_has_dots else WD_TAB_LEADER.SPACES
                p.paragraph_format.tab_stops.add_tab_stop(Inches(6.3), WD_TAB_ALIGNMENT.RIGHT, leader)

        # 3. MIGRACION DE NUMEROS DE PAGINA A PIES DE PAGINA (NATIVOS)
        print(f"[INFO] Migrando numeros de pagina a pies nativos...")
        
        # Identificar los parrafos que son numeros de pagina y a que seccion pertenecen
        section_paragraphs = []
        current_group = []
        
        for p in wordoc.paragraphs:
            current_group.append(p)
            # Detectar salto de seccion en las propiedades del parrafo (XML)
            if p._p.xpath('./w:pPr/w:sectPr'):
                section_paragraphs.append(current_group)
                current_group = []
        if current_group:
            section_paragraphs.append(current_group)
        
        # Para cada seccion, buscar el numero de pagina y moverlo
        for i, paragraphs in enumerate(section_paragraphs):
            if i >= len(wordoc.sections): continue
            
            # Buscar el parrafo que solo tiene un numero (arabe o romano)
            page_p = None
            for p in reversed(paragraphs):
                txt = p.text.strip().upper() # Usar mayusculas para el match
                if re.match(r'^\d+$', txt) or re.match(r'^[IVX]+$', txt):
                    page_p = p
                    break
            
            if page_p:
                page_text = page_p.text.strip()
                section = wordoc.sections[i]
                footer = section.footer
                
                # IMPORTANTE: Desvincular de la seccion anterior para tener numeracion independiente
                footer.is_linked_to_previous = False
                
                # Limpiar cualquier cosa que tuviera el pie
                for p_old in footer.paragraphs:
                    p_old.text = ""
                
                # Crear el nuevo parrafo en el pie
                p_foot = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
                p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_foot.paragraph_format.tab_stops.clear_all()
                run_foot = p_foot.add_run(page_text)
                run_foot.font.name = 'Times New Roman'
                run_foot.font.size = Pt(12)
                
                # Eliminar el parrafo original del cuerpo
                # (Usamos XML para removerlo de forma segura)
                page_p._element.getparent().remove(page_p._element)
                print(f"[OK] Migrado {page_text} al pie de la seccion {i}")

        # 4. FORZAR TIMES NEW ROMAN EN TODO EL DOCUMENTO
        print(f"[INFO] Aplicando Times New Roman a todo el documento...")
        if 'Normal' in wordoc.styles:
            wordoc.styles['Normal'].font.name = 'Times New Roman'
        
        def aplicar_times(paragraphs):
            for p in paragraphs:
                for run in p.runs:
                    run.font.name = 'Times New Roman'
                    if not run.font.size: run.font.size = Pt(12)

        aplicar_times(wordoc.paragraphs)
        for section in wordoc.sections:
            aplicar_times(section.footer.paragraphs)
            aplicar_times(section.header.paragraphs)
        for table in wordoc.tables:
            for row in table.rows:
                for cell in row.cells:
                    aplicar_times(cell.paragraphs)

        wordoc.save(temp_docx)
        print(f"[OK] Post-procesado completo")

        # Mover archivo temporal al destino final
        try:
            if os.path.exists(docx_path):
                os.remove(docx_path)
            shutil.move(temp_docx, docx_path)
        except PermissionError:
            final_alt = docx_path.replace(".docx", "_nuevo.docx")
            shutil.move(temp_docx, final_alt)
            print(f"[WARN] Word abierto. Archivo guardado como: {final_alt}")
        
        print(f"[OK] Archivo Word generado: {docx_path}")

    except Exception as e:
        import traceback
        print(f"[WARN] Error en post-procesado: {e}")
        traceback.print_exc()

    finally:
        # Limpiar imagenes temporales
        for img_file in extracted_images:
            try:
                os.remove(img_file)
            except: pass
        if os.path.exists(temp_docx):
            try:
                os.remove(temp_docx)
            except: pass

if __name__ == "__main__":
    main_name = sys.argv[1] if len(sys.argv) > 1 else "tfg_uvigo"
    convert_pdf_to_docx(f"{main_name}.pdf", f"{main_name}.docx")
